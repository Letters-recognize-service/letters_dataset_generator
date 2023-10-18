import auxil
import gen
import change_case
import consts

import subprocess as sb
from os.path import abspath
from random import randint, choice
import re

from docx import Document
from docx.shared import Mm, Pt
import json
from pdf2jpg import pdf2jpg
from PyPDF2 import PdfReader
from loguru import logger

# Добавляет ответственных и дедлайн в задачу, если их нет
def extend_instruction(instructions, responsibles, execution_controllers, task_responsible_org, samples_dir):
	result = []
	for idx, instr in enumerate(instructions):
				result.append(dict())

				# Ответственные в новом предложении или в новом абзаце
				sep_char = ' ' if randint(0, 1) == True else '\n'

				result[idx]["task_responsibles_people"] = responsibles[idx]
				result[idx]["task_responsibles_groups"] = task_responsible_org

				resp_to_doc = change_case.create_responsible(instr, responsibles[idx])
				result[idx]["task_text"] = sep_char + resp_to_doc + '.'

				if ("оставляю за собой" in resp_to_doc):
					result[idx]["task_responsibles_people"] = "Автор приказа."


				# Дедлайн в новом предложении или в новом абзаце
				sep_char = ' ' if randint(0, 1) == True else '\n'
				deadline = auxil.generate_date(unixtime=True)

				with open(f"{samples_dir}/task_deadline.txt") as ddfile:
					dd_list = ddfile.read().split('\n')
					deadline_msg = choice(dd_list)
					
				result[idx]["task_deadline"] = deadline
				result[idx]["task_text"] += sep_char + deadline_msg
				result[idx]["task_text"] += deadline[0] + '.'

	return result

def write_docx(header, name, intro, instruction, responsible, creator,
		date, out, count, logo, sign, seal):
	document = Document()

	style = document.styles['Normal']
	font = style.font
	p_format = style.paragraph_format
	font.name = consts.font_name
	font.size = consts.font_size
	p_format.line_spacing_rule = consts.line_spacing

	for section in document.sections:
		section.top_margin = consts.top_margin
		section.bottom_margin = consts.bottom_margin
		section.left_margin = consts.left_margin
		section.right_margin = consts.right_margin

	if logo:
		document.add_picture(logo, consts.logo_w, consts.logo_h)

	# нарисовать название отдела, инициировавшего документ
	headerp = document.add_paragraph()
	headerp.alignment = 1
	headerp.add_run(header + '\n\n').bold = True
	
	# тип документа (приказ или распоряжение - для хакатона), поменять название при рефакторинге
	namep = document.add_paragraph()
	namep.alignment = 1
	namep.add_run(name)

	# вводная часть
	introp = document.add_paragraph(intro)
	introp.paragraph_format.first_line_indent = consts.first_line_indent
	
	instruction = auxil.add_numbering(instruction)

	for idx, i in enumerate(instruction):
	    document.add_paragraph(i)
	    responsiblep = document.add_paragraph(responsible[idx])
	    responsiblep.paragraph_format.first_line_indent = consts.first_line_indent

	datep = document.add_paragraph(creator+'\n')
	datep.add_run(date)
	datep.alignment = 2

	if sign:
		signp = document.add_paragraph()
		signp.alignment = 2
		signr = signp.add_run()
		signr.add_picture(sign, consts.sign_w, consts.sign_h)

	if seal:
		sealp = document.add_paragraph()
		sealp.alignment = 2
		sealr = sealp.add_run()
		sealr.add_picture(seal, consts.seal_w, consts.seal_h)

	path = f"{out}/docx/{count}.docx"
	document.save(path)
	logger.debug(path)

	return path

def write_json(instruction, responsible_arr, date, out, count):
	json_dict = {
		"Tasks": {}
	}

	for i in range(len(instruction)):
		instr = instruction[i]
		task_text = instr["task_text"].strip()
		task_responsibles_people = instr["task_responsibles_people"]
		task_responsibles_groups = instr["task_responsibles_groups"]
		task_deadline = instr["task_deadline"]

		json_dict["Tasks"][f"Task{i+1}"] = {"task_text": task_text}
		json_dict["Tasks"][f"Task{i+1}"]["task_responsibles_people"] = task_responsibles_people
		json_dict["Tasks"][f"Task{i+1}"]["task_responsibles_groups"] = task_responsibles_groups
		json_dict["Tasks"][f"Task{i+1}"]["task_deadline"] = task_deadline


	json_dict["Tasks"]["Global_supervisor"] = responsible_arr
	json_dict["Tasks"]["Global_deadline"] = date

	with open(f"{out}/json/{count}.json", "w") as jsonf:
		json.dump(json_dict, jsonf, ensure_ascii=False, indent=4)
		logger.debug(f"Saved {out}/json/{count}.json")

	return f"{out}/json/{count}.json"

def write_pdf_linux(docx_path, out, count):
	out_path = f"{out}/pdf/{count}.pdf"
	out_path = abspath(out_path)
	docx_path = abspath(docx_path)

	cmd = ""
	cmd += f"abiword "
	cmd += f"-t pdf "
	cmd += f"-o {out_path} "
	cmd += docx_path
	sb.call(cmd, shell=True, stderr=sb.DEVNULL)
	logger.debug(f"Saved {out_path}")

	return out_path

def write_jpg(out, count):
	
	pdf = f"{out}/pdf/{count}.pdf"
	output = f"{out}/jpg"

	pdf2jpg.convert_pdf2jpg(pdf, output, pages="ALL")
	logger.debug(f"Saved {out}/jpg/{count}.pdf_dir/")

def extract_tm(pdf_path, page_num):
	reader = PdfReader(pdf_path)
	page = reader.pages[page_num]

	tmx, tmy = [], []
	def visitor_sign(text, cm, tm, fontDict, fontSize):
		tmx.append(tm[4])
		tmy.append(tm[5])

	text = page.extract_text(visitor_text=visitor_sign)

	# Координаты последнего блока с текстом
	# в PdfUnits
	tmx = tmx[-1]
	tmy = tmy[-1]

	return (tmx, tmy, len(page.images), len(reader.pages) + page_num) # page_num = {-1, -2} поэтому +

def write_coords(json_path, pdf_path, data, is_image=False):

	with open(json_path, "r") as json_file:
		json_dict = json.load(json_file)

	if is_image:

		# Координаты логотипа
		logo_coords = auxil.calculate_logo_coords()

		# Координаты подписи
		(tmx, tmy, im_count, seal_page_num) = extract_tm(pdf_path, -1)

		if im_count >= 2: # Если на последней странице есть подпись и печать
			sign_page_num = seal_page_num

			if (tmx == 0) or (tmy == 0): # Если на странице только подпись и печать
				sign_coords = auxil.calculate_sign_coords(tmx, tmy, new_page=True)
				
			else:
				sign_coords = auxil.calculate_sign_coords(tmx, tmy)

		else: # Если на последней странице только печать
			(tmx, tmy, _, sign_page_num) = extract_tm(pdf_path, -2)
			sign_coords = auxil.calculate_sign_coords(tmx, tmy)

		# Координаты печати
		if im_count >= 2:
			seal_coords = auxil.calculate_seal_coords(sign_coords)

		else:
			seal_coords = auxil.calculate_seal_coords([], new_page=True)
		
		json_dict["Images"] = {}
		json_dict["Images"]["logo"] = {"coords": logo_coords,
                                       "page_num": 0}
		json_dict["Images"]["signature"] = {"coords": sign_coords,
                                            "page_num": sign_page_num}
		json_dict["Images"]["seal"] = {"coords": seal_coords,
                                                   "page_num": seal_page_num}

	text_coords = auxil.calculate_text_coords(pdf_path, data)
	json_dict["Text"] = {}
	json_dict["Text"]["header"] = text_coords[0]
	json_dict["Text"]["name"] = text_coords[1]
	json_dict["Text"]["intro"] = text_coords[2]
	json_dict["Text"]["tasks"] = text_coords[3]
	json_dict["Text"]["responsible"] = text_coords[4]
	json_dict["Text"]["creator"] = text_coords[5]
	json_dict["Text"]["date"] = text_coords[6]

	with open(json_path, "w") as jsonf:
		json.dump(json_dict, jsonf, ensure_ascii=False, indent=4)
